


from docx import Document
from docx.shared import Inches
# from docx.enum.text import WD_ALIGN_PARAGRAPH

name = "Duc Nguyen"
address = "Victoria BC"
phone = "250-884-6325"
mail = "dukeng@uvic.ca"

def format_size_and_font(paraObj):
    paraObj.font.size = 


document = Document()

document.add_heading(name, 2)

contact_info = document.add_paragraph("ERWERWERWERWERWER")

# contact_info.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

# set first table of contact info

# contact_info.add_run(address)
# contact_info.add_run(phone)
# contact_info.add_run(mail)


paras = document.paragraphs

idx = 0
for para in paras:
    print("________", idx, "____________")
    idx += 1
    print(para.text)
    para.clear()

document.save("news.docx")