from docx import Document
from docx.shared import Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt 
import datetime
import json, re
from pprint import pprint

# dynamic info
name = "Duc Nguyen"
address = "Victoria BC"
phone = "250-884-6325"
email = "dukeng@uvic.ca"
website="https://dukeng.github.io/"
github="https://github.com/dukeng/"


data = json.load(open('info.json'))


company_name=data["company_name"]
position=data["position"]
open_para=data["open_para"]
first_coop=data["first_coop"]
second_coop=data["second_coop"]
activities=data["activities"]
closing=data["closing"]

document = Document()

document.add_heading(name, 2)

margin = 2
#changing the page margins
sections = document.sections
for section in sections:
    section.top_margin = Cm(margin)
    section.bottom_margin = Cm(margin)
    section.left_margin = Cm(margin)
    section.right_margin = Cm(margin)


contact_info = document.add_paragraph()

def format_size_and_font(format_obj):
    format_obj.font.size = Pt(12)

def format_alignment(para_obj):
    para_obj.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    para_obj.paragraph_format.first_line_indent = Inches(0.5)
    return para_obj

def format_fill_in_info(format_obj):
    replaceList = re.findall("\{.*?\}", format_obj)
    for replace in replaceList:
        replaceTo = replace.strip("{").strip("}")
        format_obj = re.sub(replace, data[replaceTo], format_obj)
    return format_obj

def askForChoice(choice_obj):
    print("Choose one of the following:")
    choices = []
    for key, value in choice_obj.items():
        print(key, end=", ")
        choices.append(value)
    print(":")
    choice = input()
    return choices[int(choice)]

def askForChoices(choice_obj):
    print("Choose one or more of the following. Enter -1 to skip:")
    choices = []
    
    for key, value in choice_obj.items():
        choices.append(key)
    
    returnChoices = []
    while(1):
        for aChoice in choices:
            print(aChoice, end=", ")
        print(":")        
        index = int(input())
        if index == -1:
            break
        returnChoices.append(choices.pop(index))
    
    returnChoices = [choice_obj[choice] for choice in returnChoices]
    return returnChoices
    


#Contact information
address_obj = contact_info.add_run(address + " , ")
phone_obj = contact_info.add_run(phone + " , ")
mail_obj = contact_info.add_run(email)
website_obj = contact_info.add_run("\n" + "Website:" + website)
github_obj = contact_info.add_run("\n" + "Github:" + github)

document.add_heading("_" * 90, 6)

# datetime
today = datetime.datetime.today()
date_obj =  document.add_paragraph().add_run(today.strftime('%d, %b %Y'))


re_obj = document.add_paragraph()
re_obj.add_run("Re: " + position)

dear_obj = document.add_paragraph()
dear_obj.add_run("Dear Hiring Manager,")


#Open paragraph
openpara_obj = document.add_paragraph()
openpara_obj = format_alignment(openpara_obj)

open_para_choice = askForChoice(open_para)

openpara_obj.add_run(format_fill_in_info(open_para_choice))


#First coop
first_coop_obj = document.add_paragraph()
first_coop_obj = format_alignment(first_coop_obj)
first_coop_obj.add_run(first_coop)

#Second coop
second_coop_obj = document.add_paragraph()
second_coop_obj.add_run(second_coop)
second_coop_obj = format_alignment(second_coop_obj)


#Activies
activities_obj = document.add_paragraph()
activities_obj = format_alignment(activities_obj)


activities = askForChoices(activities)

for activity in activities:
    activities_obj.add_run(activity + ". ")

#Final paragraph
final_obj = document.add_paragraph()
final_obj = format_alignment(final_obj)
final_obj.add_run(format_fill_in_info(closing))

#Closing
closing_obj= document.add_paragraph()
closing_obj = format_alignment(closing_obj)
closing_obj.add_run("Sincerely,")

#Signature
name_obj= document.add_paragraph()
name_obj = format_alignment(name_obj)
name_obj.add_run(name)

# document.add_page_break()

document.save('news.docx')